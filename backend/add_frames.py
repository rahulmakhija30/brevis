from youtube_transcript_api import YouTubeTranscriptApi
import docx
from docx.shared import Inches
import os
from bs4 import BeautifulSoup as bs
import requests
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from web_scraping import web_scrape

directory = os.getcwd()+'/out'

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def get_title(url):
        r = requests.get(url) 
        s = bs(r.text, "html.parser") 
        title = s.find("span", class_="watch-title").text.replace("\n", "") 
        return title.strip()

def add_picture(url,scrape_results):
        urlID = url.partition('https://www.youtube.com/watch?v=')[-1]
        transcript = YouTubeTranscriptApi.get_transcript(urlID)
        vid_title=get_title(url)
        vid_title="Notes on " + vid_title



        f1=open('summary.txt')
        summary=f1.read()

        f1.close()

        #f=open("file1.txt","w")
        s=set()
        paras=[[i,[]]for i in summary.split('\n')]
        j=0
        for filename in os.listdir(directory):
                l=filename.split(".")
                time=float(l[0][5:])
                j=0
                while(j<len(transcript)):
                        data=transcript[j]
                        index=0
                        if(time>=(data['start']*1000) and time<(((data['start']+data['duration'])*1000)+2000)):
                                text=data['text'].replace('\n',' ')
                                t=(text,time)
                                while index <(len(paras)):
                                        if(text in paras[index][0]):
                                                if(filename not in paras[index][1]):
                                                        paras[index][1].append(filename)
                                                break
			
                                        index+=1
                        j+=1

        document = docx.Document()
        d=document.add_heading(vid_title,0)
        d.alignment=1

        l=document.add_heading("", level=2)
        l.alignment=1
        l.add_run("Link :  " + str(url)).underline=True
        
        c=document.add_heading('', level=1)
        c.alignment=0
        c.add_run("Content : ").underline=True
        
        

        p = document.add_paragraph()
        r = p.add_run()
        r.add_break()
        img_no=1
        for para in paras:
                if(para[0]):
                        #f.write(para[0])
                        #f.write("\n\n")
                        r.add_text(para[0])
                        #f.write("\n\n")
                        r.add_break()
                        r.add_break()
                if(para[1]):
                        for name in para[1]:
                                p=document.add_paragraph()
                                r=p.add_run()
                                r.add_picture(directory+'/'+name,width=Inches(5.0))
                                r.add_break()
                                p.add_run("Fig. "+str(img_no)).italics=True
                                img_no+=1
                                paragraph = document.paragraphs[-1]
                                paragraph.alignment = 1
                                p=document.add_paragraph()
                                r=p.add_run()
                                r.add_break()
                                #f.write("\n\n")
                                #f.write(name)
                                #f.write("\n\n")
                                r.add_break()
                                r.add_break()
        #Adding Scraped links to the file 
        l=document.add_heading('', level=1)
        l.alignment=0
        l.add_run("External References : ").underline=True
        l.add_run().add_break()
        p=document.add_paragraph()
        p.add_run("For more details regarding the topic, please visit :")
        google_links=scrape_results["google"]
        youtube_links=scrape_results["youtube"]
        #add_hyperlink(p,"https://www.google.com","Google","0000FF", True)
        h=document.add_heading("",level=2)
        h.add_run("Google References :").underline=True
        p=document.add_paragraph()
        p.add_run().add_break()
        for d in google_links:
                add_hyperlink(p,d["linktopage"],d["title"],"0000FF",True)
                p.add_run().add_break()
        h=document.add_heading("",level=2)
        h.add_run("Youtube References :").underline=True
        p=document.add_paragraph()
        p.add_run().add_break()
        for d in youtube_links:
                add_hyperlink(p,d["linktopage"],d["title"],"0000FF",True)
                p.add_run().add_break()
        
        document.save('test1.docx')
        #f.close()

if __name__ == "__main__":
    s=web_scrape(["Artificial Intelligence"])
    add_picture("https://www.youtube.com/watch?v=ukzFI9rgwfU",s)

