# !pip install youtube_transcript_api
# !pip install python-docx

from youtube_transcript_api import YouTubeTranscriptApi
from docx.shared import Inches
from bs4 import BeautifulSoup as bs
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from web_scraping import Scrapper

from gensim.models.wrappers import FastText
from gensim.models import FastText
from fse.models import Average
from fse import IndexedList

import os
import docx
import requests
import shutil
import re
import pysbd
import time
import logging
import pafy

import platform
import os
from docx2pdf import convert

logging.basicConfig(format='%(asctime)s : %(threadName)s : %(levelname)s : %(message)s',level=logging.INFO)

import warnings
warnings.filterwarnings("ignore")

directory = os.path.join('res','out')

class Notes:
	def __init__(self,url,scrape_results):
		self.url = url
		self.scrape_results = scrape_results
		
		if self.scrape_results == {}:
			self.SCRAPE_STATUS = 0
		else:
			self.SCRAPE_STATUS = 1

	def clean(self,text):
		clean = re.sub(r"""
				   [,.;@#?!&$]+  # Accept one or more copies of punctuation
				   \ *           # plus zero or more copies of a space,
				   """,
				   " ",          # and replace it with a single space
				   text, flags=re.VERBOSE).lower().strip()
		return clean

	def modify(self,string):
		return re.sub('[^A-Za-z0-9]+', ' ', string).strip().lower().split()

	def sentence_similarity(self,text1,text2,similarity_threshold=0.35):
		sentences = [self.modify(text1),self.modify(text2)]
		ft = FastText(sentences, min_count=1,size=12,workers=4)

		model = Average(ft)
	
		try:
			model.train(IndexedList(sentences),update=True,report_delay=10,queue_factor=4)
		except ZeroDivisionError as z:
			pass

		sim = model.sv.similarity(0,1)

		if sim >= similarity_threshold:
			return True

		else:
			return False

	def add_hyperlink(self, paragraph, text, place_url, color, underline):
		"""
		A function that places a hyperlink within a paragraph object.

		:param paragraph: The paragraph we are adding the hyperlink to.
		:param place_url: A string containing the required url
		:param text: The text displayed for the url
		:return: The hyperlink object
		"""

		# This gets access to the document.xml.rels file and gets a new relation id value
		part = paragraph.part
		r_id = part.relate_to(place_url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

		# Create the w:hyperlink tag and add needed values
		hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
		hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

		# Create a w:r element
		new_run = docx.oxml.shared.OxmlElement('w:r')

		# Create a new w:rPr element
		rPr = docx.oxml.shared.OxmlElement('w:rPr')

		# Add color if it is given
		if not color is None:
			c = docx.oxml.shared.OxmlElement('w:color')
			c.set(docx.oxml.shared.qn('w:val'), color)
			rPr.append(c)

		# Remove underlining if it is requested
		if not underline:
			u = docx.oxml.shared.OxmlElement('w:u')
			u.set(docx.oxml.shared.qn('w:val'), 'none')
			rPr.append(u)

		# Join all the xml elements together add add the required text to the w:r element
		new_run.append(rPr)
		new_run.text = text
		hyperlink.append(new_run)

		paragraph._p.append(hyperlink)

		return hyperlink

	'''
	def get_title(url):
		r = requests.get(url) 
		s = bs(r.text, "html.parser") 
		title = s.find("span", class_="watch-title")

		while title == None:
			r = requests.get(url) 
			s = bs(r.text, "html.parser") 
			title = s.find("span", class_="watch-title")

		return title.text.replace("\n", "") .strip()
	'''

	def get_title(self):
		video = pafy.new(self.url)
		self.channel_name=video.author
		return video.title
		

	def generate_notes(self):
		urlID = self.url.partition('https://www.youtube.com/watch?v=')[-1]
		transcript = YouTubeTranscriptApi.get_transcript(urlID)
		vid_title = self.get_title()
		vid_title="Notes on " + vid_title

		f1=open(os.path.join('res','paragraph_headings.txt'),encoding="utf-8")
		summary=f1.read()

		f1.close()

		# Extracting Heading and Paragraphs
		data = summary.split('\n')
		heading = []
		para = ""
		for i in range(len(data)-1):
			temp = data[i].split('$')
			heading.append(temp[0])
			para = para + temp[1] + "\n"

		#f=open("file1.txt","w")
		s=set()
		paras=[[i,[]]for i in para.split('\n')]
		paras = paras[:-1]
		j=0
		for filename in os.listdir(directory):
			l=filename.split(".")
			time=float(l[0][5:])
			j=0
			
			while(j<len(transcript)):
				data=transcript[j]
				index=0
				
				if(time>=(data['start']*1000) and time<(((data['start']+data['duration'])*1000)+2000)):
					text=data['text'].replace('\n',' ')
					t=(text,time)
					
					while index <(len(paras)):
						t1 = self.clean(text)
						t2 = self.clean(paras[index][0])
						
						if((t1 in t2) or self.sentence_similarity(t1,t2)):
							if(filename not in paras[index][1]):
								paras[index][1].append(filename)
							break
							
						index+=1
						
				j+=1
				

		document = docx.Document()
		d=document.add_heading(vid_title,0)
		d.alignment=1

		l=document.add_heading("", level=2)
		l.alignment=1
		l.add_run("Link :  " + str(self.url)).underline=True

		c=document.add_heading('', level=1)
		c.alignment=0
		c.add_run("Content : ").underline=True


		p = document.add_paragraph()
		r = p.add_run()
		r.add_break()
		img_no=1
		e=0
		# print(len(paras))
		for para in paras:
			if(para[0]):
				#f.write(para[0])
				#f.write("\n\n")
				h=document.add_heading("",level=2)
				h.alignment=0
				h.add_run(heading[e]).underline = True
				h.add_run().add_break()
				p = document.add_paragraph()
				r = p.add_run()
				r.add_text(para[0])
				#f.write("\n\n")
				r.add_break()
				r.add_break()
				e+=1
			if(para[1]):
				for name in para[1]:
					p=document.add_paragraph()
					r=p.add_run()
					r.add_picture(directory+'/'+name,width=Inches(5.0))
					r.add_break()
					p.add_run("Fig. "+str(img_no)).italics=True
					img_no+=1
					paragraph = document.paragraphs[-1]
					paragraph.alignment = 1
					p=document.add_paragraph()
					r=p.add_run()
					r.add_break()
					#f.write("\n\n")
					#f.write(name)
					#f.write("\n\n")
					r.add_break()
					r.add_break()

		if self.SCRAPE_STATUS == 1:
			#Adding Scraped links to the file 
			l=document.add_heading('', level=1)
			l.alignment=0
			l.add_run("External References : ").underline=True
			l.add_run().add_break()
			p=document.add_paragraph()
			p.add_run("For more details regarding the topic, please visit :")
			google_links = self.scrape_results["google"]
			youtube_links = self.scrape_results["youtube"]
			#add_hyperlink(p,"https://www.google.com","Google","0000FF", True)
			h=document.add_heading("",level=2)
			h.add_run("Google References :").underline=True
			p=document.add_paragraph()
			p.add_run().add_break()
			link_no=1
			for d in google_links:
				p.add_run(str(link_no)+".> ")
				link_no+=1
				self.add_hyperlink(p,d["title"],d["linktopage"],"0000FF",True)
				p.add_run().add_break()
			h=document.add_heading("",level=2)
			h.add_run("Youtube References :").underline=True
			p=document.add_paragraph()
			p.add_run().add_break()
			link_no=1
			for d in youtube_links:
				p.add_run(str(link_no)+".> ")
				link_no+=1
				self.add_hyperlink(p,d["title"],d["linktopage"],"0000FF",True)
				p.add_run().add_break()
						

		#Adding copyright notice to notes
		h=document.add_heading("",level=2)
		h.add_run("Disclaimer :").underline=True
		p=document.add_paragraph()
		p.add_run().add_break()
		p.add_run("The contents of these notes are solely derived from the video. The contents of this video is created by "+str(self.channel_name)+" channel. We are not responsible for providing any irrelevant content or errors that might have crept in. The information contained in these notes is intended to be used only for educational purposes & no part of it should be reproduced in any form without prior permission from us.")
		p.add_run().add_break()

		#Removing temporary files
		#if os.path.exists('res'):
		#    shutil.rmtree('res')
			
		document.save(os.path.join('res','Brevis-Notes.docx'))
		# document.save('Brevis-Notes.docx')

		if platform.system() == "Windows":
			word_path = os.path.join(os.path.abspath(os.getcwd()),"Brevis-Notes.docx")
			pdf_path_test = os.path.join(os.path.abspath(os.getcwd()),"Brevis-Notes.pdf")
			pdf_path = os.path.join(os.path.abspath(os.getcwd()),"res","Brevis-Notes.pdf")
			convert(word_path,pdf_path)

		# os.remove("video.mp4")
		#f.close()

if __name__ == "__main__":
	keywords = ['open source mobile sdk', 'lightweight object designed', 'experienced flutter developer', 'flutter youtube channel', 'important flutter widgets', 'important step', 'important features', 'important properties', 'apps made', 'specific task']
	
# 	scraped_results = Scrapper(keywords,2,2,2)
# 	scraped_results.web_scrape()
# 	s = scraped_results.scrape_result
	
	s = {}
	url = "https://www.youtube.com/watch?v=b_sQ9bMltGU"
	notes = Notes(url,s)
	notes.generate_notes()
	print("Brevis-Notes.docx Generated")
